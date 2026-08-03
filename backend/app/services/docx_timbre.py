"""Aplica cabecalho e rodape no .docx com python-docx.

Regra que nao pode ser quebrada: as linhas montadas aqui precisam ser
IDENTICAS as do preview do frontend (`linhasEscritorio` em preview-a4.tsx).
Se as duas logicas divergirem, o preview passa a mentir — e o preview e o
argumento de venda da tela 4.4.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

from app.models import ConfiguracaoTenant, Escritorio

# cinza do fio do rodape: equivale ao rgba(15,23,41,.2) do preview sobre branco
SEPARADOR_COR = "CFD2D8"

ALINHAMENTOS = {
    "esquerda": WD_ALIGN_PARAGRAPH.LEFT,
    "centro": WD_ALIGN_PARAGRAPH.CENTER,
    "direita": WD_ALIGN_PARAGRAPH.RIGHT,
}


def linhas_escritorio(escritorio: Escritorio | None, campos: dict) -> list[str]:
    """Espelho de `linhasEscritorio` do frontend. Manter as duas em sincronia.

    Usada tanto pelo cabecalho quanto pelo rodape — os dois escolhem de que
    campos do escritorio querem, de forma independente.
    """
    if escritorio is None:
        return []

    linhas: list[str] = []

    if campos.get("razao_social"):
        linhas.append(escritorio.razao_social)

    identidade: list[str] = []
    if campos.get("cnpj") and escritorio.cnpj:
        identidade.append(f"CNPJ {escritorio.cnpj}")
    if campos.get("oab") and escritorio.oab_numero:
        seccional = f"/{escritorio.oab_seccional}" if escritorio.oab_seccional else ""
        identidade.append(f"OAB {escritorio.oab_numero}{seccional}")
    if identidade:
        linhas.append(" · ".join(identidade))

    if campos.get("endereco"):
        rua = ", ".join(p for p in (escritorio.logradouro, escritorio.numero) if p)
        local = " - ".join(
            p for p in (escritorio.bairro, escritorio.cidade, escritorio.uf) if p
        )
        endereco = " · ".join(p for p in (rua, local, escritorio.cep) if p)
        if endereco:
            linhas.append(endereco)

    contato: list[str] = []
    if campos.get("telefone") and escritorio.telefone:
        contato.append(escritorio.telefone)
    if campos.get("whatsapp") and escritorio.whatsapp:
        contato.append(f"WhatsApp {escritorio.whatsapp}")
    if campos.get("email") and escritorio.email:
        contato.append(escritorio.email)
    if campos.get("site") and escritorio.site:
        contato.append(escritorio.site)
    if contato:
        linhas.append(" · ".join(contato))

    return linhas


def _aplicar_fonte(run, tipografia: dict, negrito_forcado: bool | None = None) -> None:
    fonte = tipografia.get("fonte", "Arial")
    run.font.size = Pt(tipografia.get("tamanho", 10))
    run.font.bold = tipografia.get("negrito", False) if negrito_forcado is None else negrito_forcado
    run.font.italic = tipografia.get("italico", False)

    cor = tipografia.get("cor", "#1E2D6B").lstrip("#")
    run.font.color.rgb = RGBColor.from_string(cor.upper())

    # python-docx so escreve o nome em w:ascii. Sem preencher hAnsi/cs/eastAsia
    # o Word e o LibreOffice caem na fonte do tema em parte do texto.
    run.font.name = fonte
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for atributo in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(atributo), fonte)


def _campo(paragrafo: Paragraph, instrucao: str, tipografia: dict) -> None:
    """Insere um campo do Word (PAGE / NUMPAGES).

    python-docx nao expoe campos: e preciso montar o XML na mao, senao a
    numeracao ficaria congelada no numero da primeira pagina.
    """
    run = paragrafo.add_run()
    _aplicar_fonte(run, tipografia)

    inicio = OxmlElement("w:fldChar")
    inicio.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instrucao} "

    fim = OxmlElement("w:fldChar")
    fim.set(qn("w:fldCharType"), "end")

    run._element.append(inicio)
    run._element.append(instr)
    run._element.append(fim)


def _texto(paragrafo: Paragraph, texto: str, tipografia: dict, negrito: bool | None = None) -> None:
    run = paragrafo.add_run(texto)
    _aplicar_fonte(run, tipografia, negrito)


def _limpar(container) -> None:
    """Remove o conteudo pre-existente do cabecalho/rodape do template."""
    for paragrafo in list(container.paragraphs):
        p = paragrafo._element
        p.getparent().remove(p)
    for tabela in list(container.tables):
        t = tabela._element
        t.getparent().remove(t)


def _sem_bordas(tabela) -> None:
    tbl_pr = tabela._element.tblPr
    borders = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    tbl_pr.append(borders)


def _separador(container) -> None:
    """Fio horizontal acima do rodape, com respiro em cima e embaixo.

    Espelha o `borderTop` do preview (`preview-a4.tsx`): 0.75pt — no OOXML a
    espessura vem em oitavos de ponto, logo `w:sz=6`. E um paragrafo vazio de
    fonte minima cuja borda inferior desenha o fio; assim a linha nasce
    colada no topo do bloco de texto, sem depender de tabela nem de shape.
    """
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = Pt(1)

    p_pr = p._element.get_or_add_pPr()
    bordas = OxmlElement("w:pBdr")
    fio = OxmlElement("w:bottom")
    fio.set(qn("w:val"), "single")
    fio.set(qn("w:sz"), "6")
    fio.set(qn("w:space"), "0")
    fio.set(qn("w:color"), SEPARADOR_COR)
    bordas.append(fio)
    p_pr.append(bordas)

    # o paragrafo nao tem texto; a fonte de 1pt evita que ele empurre o
    # rodape para dentro da area util da pagina
    run = p.add_run()
    run.font.size = Pt(1)


def aplicar(
    origem: Path,
    destino: Path,
    config: ConfiguracaoTenant,
    escritorio: Escritorio | None,
    logo: Path | None,
) -> None:
    """Le o .docx original, injeta cabecalho e rodape, salva em `destino`.

    O arquivo de origem NUNCA e modificado — e o anexo do cliente.
    """
    doc = Document(str(origem))

    tip_cab = dict(config.cabecalho_tipografia or {})
    tip_rod = dict(config.rodape_tipografia or {})
    linhas = linhas_escritorio(escritorio, dict(config.cabecalho_campos or {}))
    linhas_rodape = linhas_escritorio(escritorio, dict(config.rodape_campos or {}))
    posicao_logo = config.logo_posicao or "esquerda"
    usar_logo = logo is not None and logo.exists() and posicao_logo != "sem_logo"

    numerar = bool(config.rodape_numeracao)
    lado_num = config.rodape_numeracao_alinhamento or "direita"
    numerar_no_cabecalho = numerar and (config.numeracao_local or "rodape") == "cabecalho"

    for secao in doc.sections:
        # ---------------- cabecalho ----------------
        cabecalho = secao.header
        cabecalho.is_linked_to_previous = False
        _limpar(cabecalho)

        if numerar_no_cabecalho:
            # numeracao antes do timbre, na primeira linha do cabecalho
            p = cabecalho.add_paragraph()
            p.alignment = ALINHAMENTOS.get(lado_num)
            _numeracao(p, tip_rod)

        if usar_logo and posicao_logo in ("esquerda", "direita"):
            tabela = cabecalho.add_table(rows=1, cols=2, width=secao.page_width)
            tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
            _sem_bordas(tabela)

            celula_logo, celula_texto = (
                (tabela.cell(0, 0), tabela.cell(0, 1))
                if posicao_logo == "esquerda"
                else (tabela.cell(0, 1), tabela.cell(0, 0))
            )

            p_logo = celula_logo.paragraphs[0]
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(str(logo), height=Pt(34))

            primeiro = True
            alvo = celula_texto.paragraphs[0]
            for linha in linhas:
                p = alvo if primeiro else celula_texto.add_paragraph()
                p.alignment = ALINHAMENTOS.get(tip_cab.get("alinhamento", "centro"))
                _texto(p, linha, tip_cab, negrito=True if primeiro else None)
                primeiro = False
        else:
            if usar_logo and posicao_logo in ("acima", "centro"):
                p = cabecalho.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(logo), height=Pt(34))

            # em "centro" o texto acompanha o logo, ignorando o alinhamento
            # escolhido na tipografia — e o que faz o bloco parecer timbrado
            alinhamento_texto = (
                WD_ALIGN_PARAGRAPH.CENTER
                if posicao_logo == "centro"
                else ALINHAMENTOS.get(tip_cab.get("alinhamento", "centro"))
            )
            for i, linha in enumerate(linhas):
                p = cabecalho.add_paragraph()
                p.alignment = alinhamento_texto
                _texto(p, linha, tip_cab, negrito=True if i == 0 else None)

        # ---------------- rodape ----------------
        rodape = secao.footer
        rodape.is_linked_to_previous = False
        _limpar(rodape)

        texto_rodape = (config.rodape_texto or "").strip()
        # se a numeracao subiu para o cabecalho, o rodape nao a repete
        numerar_aqui = numerar and not numerar_no_cabecalho
        lado = lado_num

        # os dados do escritorio vem primeiro, depois o texto livre
        conteudo = [*linhas_rodape]
        if texto_rodape:
            conteudo.append(texto_rodape)

        # rodape vazio nao ganha fio: uma linha solta no pe da pagina, sem
        # nada embaixo, so sujaria o documento
        if conteudo or numerar_aqui:
            _separador(rodape)

        uma_linha_so = len(conteudo) == 1
        if numerar_aqui and lado in ("esquerda", "direita") and uma_linha_so:
            # com uma unica linha da para por texto e numeracao em lados
            # opostos, na mesma altura: tabela de 2 colunas sem bordas
            tabela = rodape.add_table(rows=1, cols=2, width=secao.page_width)
            _sem_bordas(tabela)
            celula_num, celula_txt = (
                (tabela.cell(0, 0), tabela.cell(0, 1))
                if lado == "esquerda"
                else (tabela.cell(0, 1), tabela.cell(0, 0))
            )

            p_txt = celula_txt.paragraphs[0]
            p_txt.alignment = ALINHAMENTOS.get(tip_rod.get("alinhamento", "centro"))
            _texto(p_txt, conteudo[0], tip_rod)

            p_num = celula_num.paragraphs[0]
            p_num.alignment = ALINHAMENTOS.get(lado)
            _numeracao(p_num, tip_rod)
        else:
            # varias linhas: empilha e coloca a numeracao na ultima
            for linha in conteudo:
                p = rodape.add_paragraph()
                p.alignment = ALINHAMENTOS.get(tip_rod.get("alinhamento", "centro"))
                _texto(p, linha, tip_rod)
            if numerar_aqui:
                p = rodape.add_paragraph()
                p.alignment = ALINHAMENTOS.get(lado)
                _numeracao(p, tip_rod)

    doc.save(str(destino))


def _numeracao(paragrafo: Paragraph, tipografia: dict) -> None:
    _texto(paragrafo, "Pagina ", tipografia)
    _campo(paragrafo, "PAGE", tipografia)
    _texto(paragrafo, " de ", tipografia)
    _campo(paragrafo, "NUMPAGES", tipografia)
