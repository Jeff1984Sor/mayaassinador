"""Gera o PDF de teste da tela de configuracoes.

Roda o pipeline REAL (python-docx -> LibreOffice -> carimbos) sobre um
.docx de exemplo criado na hora. E a prova de que o preview da tela e o
documento final sao a mesma coisa.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.core.storage import caminho_absoluto, dir_config
from app.models import ConfiguracaoTenant, Escritorio
from app.services import conversao, docx_timbre, pdf_carimbo

PETICAO = [
    ("EXCELENTISSIMO SENHOR DOUTOR JUIZ DE DIREITO DA VARA CIVEL DA COMARCA DE SAO PAULO", True),
    (
        "FULANO DE TAL, brasileiro, casado, engenheiro, portador da cedula de identidade "
        "RG n. 00.000.000-0, inscrito no CPF sob o n. 000.000.000-00, residente e "
        "domiciliado nesta Capital, vem, respeitosamente, a presenca de Vossa Excelencia, "
        "por intermedio de seu advogado que esta subscreve, propor a presente",
        False,
    ),
    ("ACAO DE COBRANCA", True),
    (
        "em face de BELTRANO DE TAL, pelos fatos e fundamentos juridicos a seguir expostos.",
        False,
    ),
    (
        "I - DOS FATOS. As partes celebraram instrumento particular de prestacao de "
        "servicos, por meio do qual o requerente se obrigou a executar os servicos ali "
        "descritos, mediante contraprestacao pecuniaria ajustada entre as partes. "
        "Ocorre que, apesar de cumprida integralmente a obrigacao pelo requerente, o "
        "requerido deixou de efetuar o pagamento na data aprazada.",
        False,
    ),
    (
        "II - DO DIREITO. Dispoe o Codigo Civil que aquele que, por acao ou omissao "
        "voluntaria, violar direito e causar dano a outrem, fica obrigado a repara-lo. "
        "A mora do devedor e incontroversa, restando caracterizado o inadimplemento "
        "contratual.",
        False,
    ),
    (
        "III - DOS PEDIDOS. Ante o exposto, requer-se a citacao do requerido para, "
        "querendo, apresentar contestacao, sob pena de revelia, julgando-se ao final "
        "procedente o pedido para condena-lo ao pagamento do valor devido, devidamente "
        "corrigido, alem das custas processuais e honorarios advocaticios.",
        False,
    ),
    ("Nestes termos, pede deferimento.", False),
    ("Sao Paulo, na data da assinatura eletronica.", False),
]

# marcador do fecho: e aqui que a assinatura ancorada deve cair no teste
FECHO_PADRAO = "Advogado(a)"


def _docx_exemplo(destino: Path, signatario: str | None) -> Path:
    doc = Document()
    # o fecho leva o nome real do signatario para que o modo "ancora" possa
    # ser testado de verdade no PDF de exemplo
    conteudo = [
        *PETICAO,
        (signatario or FECHO_PADRAO, True),
    ]
    for texto, centralizado in conteudo:
        p = doc.add_paragraph()
        p.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if centralizado else WD_ALIGN_PARAGRAPH.JUSTIFY
        )
        run = p.add_run(texto)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = centralizado
        p.paragraph_format.space_after = Pt(10)
    doc.save(str(destino))
    return destino


def gerar(config: ConfiguracaoTenant, escritorio: Escritorio | None, slug: str) -> Path:
    """Devolve o caminho do PDF de teste recem-gerado."""
    pasta = dir_config(slug)
    exemplo = pasta / "teste_exemplo.docx"
    timbrado = pasta / "teste_timbrado.docx"
    convertido = pasta / "teste_convertido.pdf"
    final = pasta / "teste.pdf"

    _docx_exemplo(exemplo, escritorio.signatario_nome if escritorio else None)

    logo = pasta / "logo.png"
    docx_timbre.aplicar(
        origem=exemplo,
        destino=timbrado,
        config=config,
        escritorio=escritorio,
        logo=logo if logo.exists() else None,
    )

    conversao.docx_para_pdf(timbrado, convertido)

    def imagem(campo: str) -> Path | None:
        relativo = getattr(config, campo, None)
        if not relativo:
            return None
        try:
            caminho = caminho_absoluto(relativo)
        except ValueError:
            return None
        return caminho if caminho.exists() else None

    pdf_carimbo.carimbar(
        origem=convertido,
        destino=final,
        rubrica=imagem("rubrica_path"),
        assinatura=imagem("assinatura_path"),
        url_verificacao=(
            "https://exemplo/verificar/TESTE123" if config.qrcode_ativo else None
        ),
        codigo="TEST-E123" if config.qrcode_ativo else None,
        posicao=pdf_carimbo.PosicaoAssinatura(
            modo=config.assinatura_modo or "fixa",
            texto=(config.assinatura_ancora or "").strip()
            or (escritorio.signatario_nome if escritorio else None),
            relativa=config.assinatura_relativa or "abaixo",
            deslocamento=config.assinatura_deslocamento or 6,
        ),
    )

    for temporario in (exemplo, timbrado, convertido):
        temporario.unlink(missing_ok=True)

    return final
